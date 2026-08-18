"""
tailor_resume.py
For each job in matched_jobs.json, asks Gemini to rewrite the resume's
summary + bullet points to emphasize what that specific posting wants,
WITHOUT inventing new experience. Writes one .docx per job.

Requires: pip install requests python-docx
Requires env var: GEMINI_API_KEY (free - get one at https://aistudio.google.com/apikey)

Uses Gemini's free tier (gemini-2.0-flash). Free tier has a daily request
cap (currently generous enough for this - well under it even tailoring
10 resumes/day) and Google may use free-tier traffic to improve their
models, unlike paid API traffic. If that trade-off bothers you, see the
README for how to switch back to a paid model instead.
"""

import json
import os
import time
from pathlib import Path

import requests
from docx import Document
from docx.shared import Pt

DATA_DIR = Path(__file__).parent.parent / "data"
OUT_DIR = Path(__file__).parent.parent / "tailored_resumes"

GEMINI_MODEL = "gemini-2.0-flash"
GEMINI_URL = (
    f"https://generativelanguage.googleapis.com/v1beta/models/"
    f"{GEMINI_MODEL}:generateContent"
)

SYSTEM_PROMPT = """You tailor resumes. You NEVER invent skills, employers, dates, \
degrees, or achievements that are not present in the candidate's original resume \
profile. You may only: reorder bullets, rephrase bullets to mirror the job \
description's language, and rewrite the summary to foreground the most relevant \
existing skills. Output strict JSON matching the schema you're given, nothing else - \
no markdown fences, no commentary."""

USER_PROMPT_TEMPLATE = """CANDIDATE RESUME PROFILE (source of truth - do not add anything not in here):
{profile_json}

JOB POSTING:
Title: {job_title}
Company: {company}
Description:
{job_description}

Return ONLY valid JSON with this exact schema:
{{
  "summary": "2-3 sentence tailored summary",
  "highlighted_skills": ["skill1", "skill2", ...],
  "experience_bullets": ["rewritten bullet 1", "rewritten bullet 2", ...],
  "project_bullets": ["rewritten bullet 1", "rewritten bullet 2", ...],
  "why_matched": "1 sentence explaining the fit, for your own review"
}}
"""


def call_gemini(api_key, profile, job, retries=3):
    prompt = USER_PROMPT_TEMPLATE.format(
        profile_json=json.dumps(profile),
        job_title=job["title"],
        company=job["company"],
        job_description=job["description"][:4000],  # keep prompt small
    )
    body = {
        "system_instruction": {"parts": [{"text": SYSTEM_PROMPT}]},
        "contents": [{"role": "user", "parts": [{"text": prompt}]}],
        "generationConfig": {
            "temperature": 0.4,
            "maxOutputTokens": 1000,
            "responseMimeType": "application/json",
        },
    }
    last_err = None
    for attempt in range(retries):
        try:
            resp = requests.post(
                GEMINI_URL,
                params={"key": api_key},
                json=body,
                timeout=30,
            )
            if resp.status_code == 429:
                # free tier rate limit - back off and retry
                wait = 15 * (attempt + 1)
                print(f"  rate limited, waiting {wait}s...")
                time.sleep(wait)
                continue
            resp.raise_for_status()
            data = resp.json()
            text = data["candidates"][0]["content"]["parts"][0]["text"].strip()
            text = text.removeprefix("```json").removeprefix("```").removesuffix("```").strip()
            return json.loads(text)
        except Exception as e:
            last_err = e
            time.sleep(2)
    raise RuntimeError(f"Gemini call failed after {retries} attempts: {last_err}")


def build_docx(profile, tailored, job, out_path):
    doc = Document()
    style = doc.styles["Normal"]
    style.font.size = Pt(10.5)

    doc.add_heading(profile["name"], level=1)
    contact = f"{profile['phone']} | {profile['email']} | {profile['linkedin']}"
    doc.add_paragraph(contact)

    doc.add_heading("Summary", level=2)
    doc.add_paragraph(tailored["summary"])

    doc.add_heading("Skills", level=2)
    doc.add_paragraph(", ".join(tailored["highlighted_skills"]))

    doc.add_heading("Experience", level=2)
    for exp in profile["experience"]:
        doc.add_paragraph(f"{exp['title']} - {exp['org']} ({exp['dates']})", style="Intense Quote")
        for b in tailored["experience_bullets"]:
            doc.add_paragraph(b, style="List Bullet")

    doc.add_heading("Projects", level=2)
    for proj in profile["projects"]:
        doc.add_paragraph(f"{proj['title']} ({proj['dates']})", style="Intense Quote")
    for b in tailored["project_bullets"]:
        doc.add_paragraph(b, style="List Bullet")

    doc.add_heading("Education", level=2)
    for edu in profile["education"]:
        doc.add_paragraph(f"{edu['degree']}, {edu['institution']} ({edu['dates']}) - {edu['score']}")

    doc.add_heading("Certifications", level=2)
    for c in profile["certifications"]:
        doc.add_paragraph(c, style="List Bullet")

    doc.save(out_path)


def main():
    profile = json.loads((DATA_DIR / "resume_profile.json").read_text())
    jobs = json.loads((DATA_DIR / "matched_jobs.json").read_text())

    api_key = os.environ["GEMINI_API_KEY"]
    OUT_DIR.mkdir(exist_ok=True)

    results = []
    for job in jobs:
        print(f"Tailoring for: {job['title']} @ {job['company']}")
        try:
            tailored = call_gemini(api_key, profile, job)
        except Exception as e:
            print(f"  failed: {e}")
            continue
        time.sleep(4)  # stay comfortably under free-tier rate limits

        safe_id = job["id"].replace("/", "-")
        docx_path = OUT_DIR / f"{safe_id}.docx"
        build_docx(profile, tailored, job, docx_path)

        job["tailored"] = tailored
        job["tailored_resume_file"] = str(docx_path.name)
        results.append(job)

    (DATA_DIR / "final_results.json").write_text(json.dumps(results, indent=2))
    print(f"Done. {len(results)} tailored resumes written to {OUT_DIR}")


if __name__ == "__main__":
    main()
